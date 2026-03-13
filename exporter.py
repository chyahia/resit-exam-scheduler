import io
import docx
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def setup_landscape_doc():
    doc = docx.Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    return doc

def make_table_rtl(table):
    """قلب اتجاه الجدول ليكون من اليمين لليسار"""
    tblPr = table._element.tblPr
    for b in tblPr.findall(qn('w:bidiVisual')): 
        tblPr.remove(b)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

def set_cell_background(cell, color_hex):
    """تلوين خلفية الخلية"""
    tcPr = cell._element.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
        
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    tcPr.append(shading_elm)

def format_paragraph(p, font_size=14, bold=False, align_center=False):
    """دالة التنسيق الآمنة (بدون اختراق XML لكي لا يرفضها الوورد)"""
    # 1. المحاذاة الرسمية الآمنة
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.RIGHT
    
    # 2. تطبيق خصائص الخط واللغة العربية
    for run in p.runs:
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'  # استخدام خط آريال لوضوحه
        if bold:
            run.font.bold = True
            
        rPr = run._element.get_or_add_rPr()
        
        # تفعيل اتجاه اليمين-لليسار للكلمات
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        
        cs = OxmlElement('w:cs')
        cs.set(qn('w:val'), '1')
        rPr.append(cs)
        
        # إجبار الحجم للخطوط العربية (Complex Scripts)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size * 2))
        rPr.append(szCs)
        
        if bold:
            bCs = OxmlElement('w:bCs')
            bCs.set(qn('w:val'), '1')
            rPr.append(bCs)

def generate_levels_word(db, distribution):
    doc = setup_landscape_doc()
    
    h = doc.add_heading('جداول الامتحانات الاستدراكية - حسب المستويات', 0)
    format_paragraph(h, font_size=20, bold=True, align_center=True)
    
    levels = db.get('levels', [])
    schedule = db.get('schedule', {})
    
    days = list(schedule.keys())
    times_list = []
    for d in days:
        for slot in schedule[d]:
            if slot['time'] not in times_list:
                times_list.append(slot['time'])
                
    for level in levels:
        h_lvl = doc.add_heading(f'المستوى: {level}', level=1)
        format_paragraph(h_lvl, font_size=16, bold=True)
        
        table = doc.add_table(rows=len(times_list) + 1, cols=len(days) + 1)
        table.style = 'Table Grid'
        make_table_rtl(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الفترة الزمنية'
        set_cell_background(hdr_cells[0], "D9E2F3") 
        for p in hdr_cells[0].paragraphs:
            format_paragraph(p, font_size=15, bold=True, align_center=True)
            
        for i, day in enumerate(days):
            hdr_cells[i+1].text = day
            set_cell_background(hdr_cells[i+1], "D9E2F3") 
            for p in hdr_cells[i+1].paragraphs:
                format_paragraph(p, font_size=15, bold=True, align_center=True)
            
        for r_idx, time_val in enumerate(times_list):
            row_cells = table.rows[r_idx + 1].cells
            
            row_cells[0].text = time_val
            set_cell_background(row_cells[0], "D9E2F3") 
            for p in row_cells[0].paragraphs:
                format_paragraph(p, font_size=14, bold=True, align_center=True)
            
            for c_idx, day in enumerate(days):
                cell = row_cells[c_idx + 1]
                day_data = distribution.get(day, {})
                if time_val in day_data:
                    levels_dict = day_data[time_val]
                    if level in levels_dict:
                        subject_name = levels_dict[level].get("subject", "مادة غير معروفة")
                        sub_teachers = levels_dict[level].get("subject_teachers", [])
                        teacher_str = "، ".join(sub_teachers) if sub_teachers else "غير محدد"
                        
                        cell_text = f"المادة: {subject_name}\nأستاذ المادة: {teacher_str}\n\nالحراسة:\n"
                        
                        for room, teachers in levels_dict[level]["rooms"].items():
                            cell_text += f"القاعة: {room}\n"
                            cell_text += "، ".join(teachers) if teachers else "بدون حراس"
                            cell_text += "\n"
                            
                        cell.text = cell_text.strip()
                
                for p in cell.paragraphs:
                    # هنا نأمر بالمحاذاة لليمين بخط واضح (14)
                    format_paragraph(p, font_size=14, bold=False, align_center=False)
        
        doc.add_page_break()
        
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem

def generate_teachers_word(db, distribution):
    doc = setup_landscape_doc()
    
    h = doc.add_heading('جداول الامتحانات الاستدراكية - للأساتذة', 0)
    format_paragraph(h, font_size=20, bold=True, align_center=True)
    
    teachers = db.get('teachers', [])
    schedule = db.get('schedule', {})
    teacher_subjects = db.get('teacher_subjects', {})
    
    days = list(schedule.keys())
    times_list = []
    for d in days:
        for slot in schedule[d]:
            if slot['time'] not in times_list:
                times_list.append(slot['time'])
                
    for teacher in teachers:
        h_t = doc.add_heading(f'الأستاذ(ة): {teacher}', level=1)
        format_paragraph(h_t, font_size=16, bold=True)
        
        subs = teacher_subjects.get(teacher, [])
        if subs:
            p = doc.add_paragraph(f"المواد المسندة: {'، '.join(subs)}")
            format_paragraph(p, font_size=14, bold=False)
        
        table = doc.add_table(rows=len(times_list) + 1, cols=len(days) + 1)
        table.style = 'Table Grid'
        make_table_rtl(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الفترة الزمنية'
        set_cell_background(hdr_cells[0], "D9E2F3")
        for p in hdr_cells[0].paragraphs:
            format_paragraph(p, font_size=15, bold=True, align_center=True)
            
        for i, day in enumerate(days):
            hdr_cells[i+1].text = day
            set_cell_background(hdr_cells[i+1], "D9E2F3") 
            for p in hdr_cells[i+1].paragraphs:
                format_paragraph(p, font_size=15, bold=True, align_center=True)
            
        for r_idx, time_val in enumerate(times_list):
            row_cells = table.rows[r_idx + 1].cells
            
            row_cells[0].text = time_val
            set_cell_background(row_cells[0], "D9E2F3")
            for p in row_cells[0].paragraphs:
                format_paragraph(p, font_size=14, bold=True, align_center=True)
            
            for c_idx, day in enumerate(days):
                cell = row_cells[c_idx + 1]
                cell_text = ""
                
                day_data = distribution.get(day, {})
                if time_val in day_data:
                    for level, level_dict in day_data[time_val].items():
                        rooms_dict = level_dict.get("rooms", {})
                        subject_name = level_dict.get("subject", "بدون مادة")
                        
                        for room, assigned_teachers in rooms_dict.items():
                            if teacher in assigned_teachers:
                                cell_text += f"المادة: {subject_name}\nالمستوى: {level}\nالقاعة: {room}\n\n"
                                
                cell.text = cell_text.strip()
                for p in cell.paragraphs:
                    # محاذاة لليمين للأساتذة أيضاً بخط كبير ومريح (14)
                    format_paragraph(p, font_size=14, bold=False, align_center=False)
                    
        doc.add_page_break()
        
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem